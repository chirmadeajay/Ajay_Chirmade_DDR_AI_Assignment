from docx import Document
from docx.shared import Pt


def write_ddr_docx(ddr_data, output_path):
    doc = Document()
    doc.styles["Normal"].font.name = "Calibri"
    doc.styles["Normal"].font.size = Pt(10.5)

    doc.add_heading("Main DDR - Detailed Diagnostic Report", level=0)
    doc.add_heading("1. Property Issue Summary", level=1)
    doc.add_paragraph(ddr_data["property_issue_summary"])

    doc.add_heading("2. Area-wise Observations", level=1)
    table = doc.add_table(rows=1, cols=4)
    table.style = "Table Grid"
    headers = ["Area", "Observation", "Severity", "Reasoning"]
    for index, header in enumerate(headers):
        table.rows[0].cells[index].text = header

    for item in ddr_data["area_wise_observations"]:
        row = table.add_row().cells
        row[0].text = item["area"]
        row[1].text = item["observation"]
        row[2].text = item["severity"]
        row[3].text = item["reasoning"]

    doc.add_heading("3. Probable Root Cause", level=1)
    doc.add_paragraph(ddr_data["probable_root_cause"])

    doc.add_heading("4. Severity Assessment", level=1)
    doc.add_paragraph("Severity is assigned based on safety risk, spread risk, and evidence strength.")

    doc.add_heading("5. Recommended Actions", level=1)
    for action in ddr_data["recommended_actions"]:
        doc.add_paragraph(action, style="List Bullet")

    doc.add_heading("6. Additional Notes", level=1)
    doc.add_paragraph("The workflow avoids unsupported assumptions and marks missing details clearly.")

    doc.add_heading("7. Missing or Unclear Information", level=1)
    for item in ddr_data["missing_or_unclear_information"]:
        doc.add_paragraph(item, style="List Bullet")

    doc.save(output_path)
